import os
import re
import json
from fpdf import FPDF
from docx import Document
from src.logger import logger

def get_base_resumes(directory="resumes"):
    if not os.path.exists(directory): os.makedirs(directory)
    return [f for f in os.listdir(directory) if f.endswith(".md") and not f.startswith("tailored/")]

def select_best_base_resume(job_title, job_description, directory="resumes"):
    resumes = get_base_resumes(directory)
    if not resumes: return None
    
    # Ensure tailored folder is ignored if it's inside resumes
    resumes = [r for r in resumes if "tailored" not in r]
    
    best_score = -1
    best_resume = resumes[0]
    
    # Extract keywords from job info
    terms = set(re.findall(r'\w+', (job_title + " " + job_description).lower()))
    
    for r_file in resumes:
        file_path = os.path.join(directory, r_file)
        content = read_resume(file_path).lower()
        # Overlap score
        score = sum(1 for term in terms if term in content)
        if score > best_score:
            best_score = score
            best_resume = file_path
            
    return best_resume

def read_resume(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f: return f.read()
    except Exception as e:
        logger.error(f"Error reading resume {file_path}: {e}")
        return ""

class ResumePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        self.set_left_margin(20)
        self.set_right_margin(20)
        self.set_font('Helvetica', '', 10)

    def header(self): pass
    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    def safe_text(self, text):
        if not text: return ""
        replacements = {'\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"', '\u2022': '*', '\u2026': '...'}
        for char, replacement in replacements.items(): text = text.replace(char, replacement)
        try: return text.encode('latin-1', 'replace').decode('latin-1')
        except: return "".join(i for i in text if ord(i) < 128)

    def render_markdown(self, content):
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: self.ln(5); continue
            self.set_x(self.l_margin)
            if line.startswith('# '):
                self.set_font("Helvetica", 'B', 16); self.multi_cell(0, 10, self.safe_text(line[2:])); self.ln(2)
            elif line.startswith('## '):
                self.ln(2); self.set_font("Helvetica", 'B', 14); self.multi_cell(0, 9, self.safe_text(line[3:])); self.ln(1)
            elif line.startswith('### '):
                self.set_font("Helvetica", 'B', 12); self.multi_cell(0, 8, self.safe_text(line[4:]))
            elif line.startswith('- ') or line.startswith('* '):
                self.set_font("Helvetica", '', 10); self.set_x(self.l_margin + 5); self.write(6, chr(149))
                self.set_x(self.l_margin + 10); eff_width = self.w - self.r_margin - (self.l_margin + 10)
                self.multi_cell(eff_width, 6, self.safe_text(line[2:]))
            else:
                self.set_font("Helvetica", '', 10); text = line.replace('**', '')
                self.multi_cell(0, 6, self.safe_text(text))

import subprocess
import shutil

def compile_pdf_with_latex(content, pdf_path):
    """
    Tries to compile the markdown resume into a professional PDF using LaTeX.
    Falls back to FPDF if LaTeX is not found or fails.
    """
    latex_cmd = None
    for cmd in ["lualatex", "xelatex", "pdflatex"]:
        if shutil.which(cmd):
            latex_cmd = cmd
            break
            
    if not latex_cmd:
        logger.info("No LaTeX compiler found (lualatex/xelatex/pdflatex). Falling back to FPDF.")
        return False

    temp_tex_path = pdf_path.replace(".pdf", ".tex")
    
    try:
        # Convert Markdown to basic LaTeX
        latex_lines = [
            r"\documentclass[11pt,a4paper]{article}",
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage[margin=0.75in]{geometry}",
            r"\usepackage{hyperref}",
            r"\usepackage{enumitem}",
            r"\setlist[itemize]{noitemsep, topsep=2pt}",
            r"\pagestyle{empty}",
            r"\begin{document}"
        ]
        
        in_list = False
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                if in_list:
                    latex_lines.append(r"\end{itemize}")
                    in_list = False
                latex_lines.append("")
                continue
                
            # Escaping special LaTeX characters safely
            clean_line = stripped.replace("&", r"\&").replace("%", r"\%").replace("$", r"\$").replace("_", r"\_")
            # Replace markdown bold **text** with \textbf{text}
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\\textbf{\1}', clean_line)
            
            if clean_line.startswith('# '):
                if in_list:
                    latex_lines.append(r"\end{itemize}")
                    in_list = False
                latex_lines.append(r"\begin{center}")
                latex_lines.append(r"{\Large \textbf{" + clean_line[2:] + r"}}")
                latex_lines.append(r"\end{center}")
            elif clean_line.startswith('## '):
                if in_list:
                    latex_lines.append(r"\end{itemize}")
                    in_list = False
                latex_lines.append(r"\section*{" + clean_line[3:] + r"}")
            elif clean_line.startswith('### '):
                if in_list:
                    latex_lines.append(r"\end{itemize}")
                    in_list = False
                latex_lines.append(r"\subsection*{" + clean_line[4:] + r"}")
            elif clean_line.startswith('- ') or clean_line.startswith('* '):
                if not in_list:
                    latex_lines.append(r"\begin{itemize}")
                    in_list = True
                latex_lines.append(r"\item " + clean_line[2:])
            else:
                if in_list:
                    latex_lines.append(r"\end{itemize}")
                    in_list = False
                latex_lines.append(clean_line + r"\\")
                
        if in_list:
            latex_lines.append(r"\end{itemize}")
            
        latex_lines.append(r"\end{document}")
        
        with open(temp_tex_path, "w", encoding="utf-8") as f:
            f.write("\n".join(latex_lines))
            
        logger.info(f"Compiling LaTeX using {latex_cmd}...")
        working_dir = os.path.dirname(pdf_path)
        subprocess.run([
            latex_cmd,
            "-interaction=nonstopmode",
            f"-output-directory={working_dir}",
            temp_tex_path
        ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=20, check=True)
        
        # Clean up auxiliary LaTeX files
        for ext in [".aux", ".log", ".out", ".tex"]:
            aux_file = pdf_path.replace(".pdf", ext)
            if os.path.exists(aux_file):
                os.remove(aux_file)
                
        logger.info(f"Successfully compiled professional LaTeX PDF: {pdf_path}")
        return True
    except Exception as e:
        logger.warning(f"LaTeX compilation failed: {e}. Falling back to FPDF.")
        # Cleanup temp tex if exists
        if os.path.exists(temp_tex_path):
            os.remove(temp_tex_path)
        return False

def save_tailored_resume(job_id, content, output_dir="resumes/tailored"):
    """Saves tailored resume as MD, PDF, and DOCX."""
    if not os.path.exists(output_dir): os.makedirs(output_dir)
    md_path = os.path.join(output_dir, f"{job_id}.md")
    pdf_path = os.path.join(output_dir, f"{job_id}.pdf")
    docx_path = os.path.join(output_dir, f"{job_id}.docx")
    
    try:
        # 1. Save MD
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(content)
            
        # 2. Save PDF (Try LaTeX, fallback to FPDF)
        compiled = compile_pdf_with_latex(content, pdf_path)
        if not compiled:
            pdf = ResumePDF()
            pdf.add_page()
            pdf.render_markdown(content)
            pdf.output(pdf_path)

        # 3. Save DOCX
        doc = Document()
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            if line.startswith('# '): doc.add_heading(line[2:], 0)
            elif line.startswith('## '): doc.add_heading(line[3:], 1)
            elif line.startswith('### '): doc.add_heading(line[4:], 2)
            elif line.startswith('- ') or line.startswith('* '): doc.add_paragraph(line[2:], style='List Bullet')
            else: doc.add_paragraph(line.replace('**', ''))
        doc.save(docx_path)
        
        return pdf_path # Return PDF as default but DOCX is now created
    except Exception as e:
        logger.error(f"Failed to save tailored resume {job_id}: {e}")
        return None
